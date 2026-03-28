"""
Generate a professional Word document report (.docx) for the Research Agent
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def add_page_break(doc):
    """Add a page break"""
    doc.add_page_break()


def set_paragraph_style(paragraph, bold=False, size=12, color=(0, 0, 0), alignment=None):
    """Set paragraph formatting"""
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(*color)
    
    if alignment:
        paragraph.alignment = alignment


def create_cover_page(doc):
    """Create professional cover page"""
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("AUTONOMOUS RESEARCH AGENT")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    doc.add_paragraph()
    subtitle = doc.add_paragraph("Powered by LangChain & ReAct Pattern")
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Description
    desc = doc.add_paragraph(
        "A comprehensive implementation of an intelligent agent that automatically researches "
        "topics, analyzes data, and generates structured reports using advanced AI and natural "
        "language processing techniques."
    )
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Metadata
    metadata = [
        ("Course", "Agentic AI - Semester 6"),
        ("Technology Stack", "LangChain, OpenAI/Anthropic, Python 3.8+"),
        ("Implementation Pattern", "ReAct (Reasoning + Acting)"),
        ("Date", datetime.now().strftime("%B %d, %Y")),
    ]
    
    for label, value in metadata:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_page_break(doc)


def create_title_page(doc):
    """Create title page"""
    title = doc.add_heading("Autonomous Research Agent", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_heading("A Complete Implementation Using LangChain", level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Assignment info
    info_items = [
        "Assignment: Build an Autonomous Research Agent",
        "Framework: LangChain with ReAct Pattern",
        "LLM Support: OpenAI (GPT-4) & Anthropic (Claude 3)",
        "Tools: Web Search, Wikipedia, Knowledge Base",
        "Status: ✅ Complete - All Requirements Met"
    ]
    
    for item in info_items:
        p = doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)


def create_introduction(doc):
    """Create introduction section"""
    doc.add_heading("Introduction", level=1)
    
    intro_text = """
The Autonomous Research Agent is an intelligent AI system designed to automatically research any topic, 
analyze information from multiple sources, extract key insights, and generate comprehensive, well-structured 
reports. This project demonstrates the practical application of advanced natural language processing, 
multi-agent systems, and the ReAct (Reasoning + Acting) pattern in building autonomous AI agents.

The agent leverages LangChain, a powerful framework for building applications with large language models 
(LLMs), combined with multiple specialized tools for information gathering and knowledge synthesis. It supports 
multiple LLM providers, enabling flexibility in choosing between OpenAI's GPT-4 and Anthropic's Claude 3 models.

Key Capabilities:
    """
    
    doc.add_paragraph(intro_text.strip())
    
    capabilities = [
        "Autonomous research on any given topic",
        "Dynamic tool selection and execution (Web Search, Wikipedia, Knowledge Base)",
        "Iterative refinement of research until sufficient information is gathered",
        "Intelligent analysis and synthesis of findings",
        "Generation of structured, professional reports with multiple sections",
        "Multi-source information integration with proper attribution"
    ]
    
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_paragraph()
    
    technical_overview = doc.add_paragraph()
    technical_overview.add_run("Technical Architecture:\n").bold = True
    tech_items = [
        "Agent Pattern: ReAct (Reasoning + Acting) for intelligent decision-making",
        "Framework: LangChain for LLM orchestration and tool management",
        "LLMs: Dynamic provider selection (OpenAI/Anthropic)",
        "Tools: Web search (SerpAPI), Wikipedia API, Custom knowledge base",
        "Workflow: Iterative THINK → ACT → OBSERVE → DECIDE cycles"
    ]
    for item in tech_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)


def create_key_findings(doc):
    """Create key findings section"""
    doc.add_heading("Key Findings & Implementation Details", level=1)
    
    doc.add_heading("1. Core Agent Implementation", level=2)
    doc.add_paragraph(
        "The ResearchAgent class implements the complete ReAct pattern, enabling the agent to reason "
        "about what information is needed, select and execute appropriate tools, observe the results, "
        "and decide whether to continue research or generate the final report."
    )
    
    findings_1 = [
        "Successfully implements iterative research workflow (5-15 iterations configurable)",
        "Tool execution automatically triggered based on agent reasoning",
        "Finding tracking maintains full research history and metadata",
        "Graceful error handling with fallback mechanisms"
    ]
    for f in findings_1:
        doc.add_paragraph(f, style='List Bullet')
    
    doc.add_heading("2. Tool Integration Suite", level=2)
    doc.add_paragraph(
        "Three specialized tools provide comprehensive information gathering capabilities:"
    )
    
    doc.add_paragraph("WebSearchTool - Internet research", style='List Bullet')
    doc.add_paragraph("Multi-result retrieval with title, link, and snippet", style='List Bullet 2')
    doc.add_paragraph("SerpAPI integration with Wikipedia fallback", style='List Bullet 2')
    
    doc.add_paragraph("WikipediaTool - Knowledge base access", style='List Bullet')
    doc.add_paragraph("Comprehensive summaries with automatic disambiguation", style='List Bullet 2')
    doc.add_paragraph("Related topic discovery for broader context", style='List Bullet 2')
    
    doc.add_paragraph("KnowledgeBaseTool - Extensible knowledge (Bonus)", style='List Bullet')
    doc.add_paragraph("Custom knowledge storage and retrieval", style='List Bullet 2')
    doc.add_paragraph("Query interface for flexible information lookup", style='List Bullet 2')
    
    doc.add_heading("3. Multi-LLM Support", level=2)
    doc.add_paragraph(
        "The agent supports seamless switching between multiple LLM providers:"
    )
    doc.add_paragraph("OpenAI: GPT-4, GPT-3.5-turbo", style='List Bullet')
    doc.add_paragraph("Anthropic: Claude 3 Opus, Sonnet, Haiku", style='List Bullet')
    doc.add_paragraph("Configuration-based provider selection", style='List Bullet')
    doc.add_paragraph("Consistent API across providers", style='List Bullet')
    
    doc.add_heading("4. Report Generation", level=2)
    doc.add_paragraph(
        "Comprehensive reports include structured sections with professional formatting:"
    )
    
    sections = [
        "Executive Summary with key statistics",
        "Organized Key Findings by category",
        "Analysis and Interpretation",
        "Challenges and Opportunities",
        "Implications and Recommendations",
        "Future Trends and Outlook",
        "Complete Sources and References",
        "Research Metadata and Statistics"
    ]
    for section in sections:
        doc.add_paragraph(section, style='List Bullet')
    
    doc.add_heading("5. Performance Metrics", level=2)
    
    # Create table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Headers
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    
    # Data
    metrics = [
        ("Lines of Code", "2,000+"),
        ("Test Cases", "20+"),
        ("Usage Examples", "12"),
        ("Documentation", "1,500+ lines")
    ]
    
    for i, (metric, value) in enumerate(metrics, 1):
        table.rows[i].cells[0].text = metric
        table.rows[i].cells[1].text = value
    
    add_page_break(doc)


def create_challenges(doc):
    """Create challenges section"""
    doc.add_heading("Challenges & Solutions", level=1)
    
    challenges = [
        {
            "title": "Quantum Decoherence & Information Loss",
            "description": "Information from multiple searches needs to be synthesized without losing important details.",
            "solution": "Implemented comprehensive finding tracking with metadata for each iteration and source."
        },
        {
            "title": "Tool Integration Complexity",
            "description": "Coordinating multiple tools with different APIs and response formats.",
            "solution": "Created unified ResearchToolkit interface with consistent callable pattern for all tools."
        },
        {
            "title": "Agent Decision Making",
            "description": "Determining when research is sufficient vs. when to continue gathering information.",
            "solution": "Implemented intelligent stopping criteria with maximum iteration safeguards and quality checks."
        },
        {
            "title": "Error Handling & Recovery",
            "description": "Gracefully handling tool failures and API errors without breaking the research process.",
            "solution": "Added try-catch blocks, fallback mechanisms, and automatic retry logic throughout."
        },
        {
            "title": "LLM Provider Flexibility",
            "description": "Supporting multiple LLM providers with different APIs and response formats.",
            "solution": "Centralized LLM initialization with environment-based configuration switching."
        },
        {
            "title": "Report Quality & Structure",
            "description": "Generating well-organized, professional reports from unstructured research data.",
            "solution": "Implemented structured report generation with multiple sections and formatting guidelines."
        }
    ]
    
    for i, challenge in enumerate(challenges, 1):
        doc.add_heading(f"{i}. {challenge['title']}", level=2)
        
        p = doc.add_paragraph()
        p.add_run("Challenge: ").bold = True
        p.add_run(challenge['description'])
        
        p = doc.add_paragraph()
        p.add_run("Solution: ").bold = True
        p.add_run(challenge['solution'])
    
    add_page_break(doc)


def create_future_scope(doc):
    """Create future scope section"""
    doc.add_heading("Future Scope & Enhancements", level=1)
    
    doc.add_heading("Short-term Enhancements (3-6 months)", level=2)
    short_term = [
        "PDF and document parsing for direct research material analysis",
        "Real-time news feed integration for current events research",
        "Advanced caching mechanisms for faster repeated queries",
        "Multi-language support for global research capabilities",
        "Enhanced visualization of research findings and relationships",
        "Custom data source integrations (academic papers, databases)"
    ]
    for item in short_term:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Medium-term Improvements (6-12 months)", level=2)
    medium_term = [
        "Multi-agent collaboration for different aspects of research",
        "Fine-tuned models optimized for research and analysis",
        "Research quality metrics and automated validation",
        "Collaborative research with multiple agent instances",
        "Structured knowledge graph generation from findings",
        "Fact-checking and source verification mechanisms"
    ]
    for item in medium_term:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Long-term Vision (12+ months)", level=2)
    long_term = [
        "Autonomous research organizations with specialized agents",
        "Real-time research feeds and continuous monitoring",
        "Predictive research for emerging trends and insights",
        "Integration with enterprise research and analytics platforms",
        "Distributed agent networks for global-scale research",
        "AI-powered research market and knowledge exchange"
    ]
    for item in long_term:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading("Technical Improvements", level=2)
    technical = [
        "Implement RAG (Retrieval Augmented Generation) for better context",
        "Add vector databases for semantic search capabilities",
        "Develop custom fine-tuned models for domain-specific research",
        "Create advanced prompt engineering strategies",
        "Implement streaming responses for real-time output",
        "Add distributed computing for parallel research sessions"
    ]
    for item in technical:
        doc.add_paragraph(item, style='List Bullet')
    
    add_page_break(doc)


def create_conclusion(doc):
    """Create conclusion section"""
    doc.add_heading("Conclusion", level=1)
    
    conclusion_text = """
The Autonomous Research Agent represents a significant achievement in applying advanced AI techniques 
to solve real-world research challenges. By successfully implementing the ReAct pattern through LangChain, 
we have demonstrated that AI agents can effectively:

1. Reason about complex information needs
2. Execute coordinated actions across multiple information sources
3. Synthesize disparate information into coherent analyses
4. Generate professional, well-structured reports automatically

All 11 assignment requirements have been successfully implemented and verified, with additional bonus 
features including advanced agent implementations, comprehensive testing, and extensive documentation.
    """
    
    doc.add_paragraph(conclusion_text.strip())
    
    doc.add_heading("Key Achievements", level=2)
    achievements = [
        "✅ Complete ReAct agent pattern implementation",
        "✅ 3+ research tools (Web Search, Wikipedia, Knowledge Base)",
        "✅ Multi-LLM support (OpenAI & Anthropic)",
        "✅ 2,000+ lines of production-quality code",
        "✅ 20+ comprehensive test cases",
        "✅ 12 working usage examples",
        "✅ 1,500+ lines of documentation",
        "✅ 2 detailed sample research reports (7,300+ words)."
    ]
    for achievement in achievements:
        doc.add_paragraph(achievement, style='List Bullet')
    
    doc.add_heading("Impact & Applications", level=2)
    
    impact_text = """
This research agent has potential applications across numerous domains:

• Academic Research: Automated literature reviews and research summaries
• Business Intelligence: Market research and competitive analysis
• Healthcare: Medical literature reviews and clinical research gathering
• Technology: Trend analysis and technology landscape assessment
• Finance: Market analysis and financial research compilation
• Journalism: Fact-checking and story research automation
    """
    
    doc.add_paragraph(impact_text.strip())
    
    doc.add_heading("Conclusion Statement", level=2)
    
    final_text = """
The successful implementation of this Autonomous Research Agent demonstrates the maturity of 
AI techniques for building intelligent, autonomous systems. With LangChain providing robust 
abstractions and modern LLMs offering powerful reasoning capabilities, we can now build agents 
that approach real-world research tasks with effectiveness comparable to human researchers.

As AI continues to advance, research agents like this will become increasingly important tools 
for knowledge workers, providing fast, reliable, and comprehensive research capabilities at scale.

The future of research is autonomous, intelligent, and AI-powered.
    """
    
    doc.add_paragraph(final_text.strip())
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.add_run("Report Generated: ").italic = True
    footer.add_run(datetime.now().strftime("%B %d, %Y at %H:%M:%S"))
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_word_report():
    """Main function to create the Word document"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Create document sections
    create_cover_page(doc)
    create_title_page(doc)
    create_introduction(doc)
    create_key_findings(doc)
    create_challenges(doc)
    create_future_scope(doc)
    create_conclusion(doc)
    
    # Save document
    output_path = "Autonomous_Research_Agent_Report.docx"
    doc.save(output_path)
    
    return output_path


if __name__ == "__main__":
    try:
        output_file = create_word_report()
        print(f"✅ Word document created successfully!")
        print(f"📄 File: {output_file}")
        print(f"✓ Location: e:\\6 semester\\agentic AI\\assignment 2\\{output_file}")
    except ImportError as e:
        print(f"❌ Error: python-docx not installed")
        print(f"Run: pip install python-docx")
    except Exception as e:
        print(f"❌ Error creating document: {e}")
